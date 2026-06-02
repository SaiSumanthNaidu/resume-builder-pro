print("STARTING APP")

from flask import Flask, render_template, request, redirect, send_file
from flask_sqlalchemy import SQLAlchemy

from reportlab.pdfgen import canvas
from io import BytesIO
from docx import Document

# CREATE FLASK APP
app = Flask(__name__)

# Database Configuration
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///resumes.db"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

db = SQLAlchemy(app)


# Resume Model
class Resume(db.Model):
    id = db.Column(db.Integer, primary_key=True)

    full_name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(120), nullable=False)
    phone = db.Column(db.String(20), nullable=False)

    summary = db.Column(db.Text)
    education = db.Column(db.Text)
    skills = db.Column(db.Text)
    experience = db.Column(db.Text)

    projects = db.Column(db.Text)
    certifications = db.Column(db.Text)
    languages = db.Column(db.Text)
    achievements = db.Column(db.Text)

    linkedin = db.Column(db.String(255))
    github = db.Column(db.String(255))
    portfolio = db.Column(db.String(255))


# Create Database Tables
with app.app_context():
    db.create_all()


# Landing Page

@app.route("/")
def landing():
    return render_template("landing.html")


# Resume Builder Page

@app.route("/builder")
def home():
    return render_template("index.html")


@app.route("/generate", methods=["POST"])
def generate():

    resume = Resume(
        full_name=request.form.get("full_name"),
        email=request.form.get("email"),
        phone=request.form.get("phone"),

        summary=request.form.get("summary"),
        education=request.form.get("education"),
        skills=request.form.get("skills"),
        experience=request.form.get("experience"),

        projects=request.form.get("projects"),
        certifications=request.form.get("certifications"),
        languages=request.form.get("languages"),
        achievements=request.form.get("achievements"),

        linkedin=request.form.get("linkedin"),
        github=request.form.get("github"),
        portfolio=request.form.get("portfolio")
    )

    db.session.add(resume)
    db.session.commit()

    return redirect(
        f"/preview/{resume.id}"
    )

@app.route("/dashboard")
def dashboard():

    search = request.args.get("search", "")

    if search:

        resumes = Resume.query.filter(
            (Resume.full_name.contains(search)) |
            (Resume.email.contains(search)) |
            (Resume.skills.contains(search))
        ).order_by(
            Resume.id.desc()
        ).all()

    else:

        resumes = Resume.query.order_by(
            Resume.id.desc()
        ).all()

    return render_template(
        "dashboard.html",
        resumes=resumes,
        search=search
    )

@app.route("/delete/<int:id>")
def delete_resume(id):

    resume = Resume.query.get_or_404(id)

    db.session.delete(resume)
    db.session.commit()

    return redirect("/dashboard")


@app.route("/edit/<int:id>", methods=["GET", "POST"])
def edit_resume(id):

    resume = Resume.query.get_or_404(id)

    if request.method == "POST":

        resume.full_name = request.form["full_name"]
        resume.email = request.form["email"]
        resume.phone = request.form["phone"]

        resume.summary = request.form["summary"]
        resume.education = request.form["education"]
        resume.skills = request.form["skills"]
        resume.experience = request.form["experience"]

        resume.projects = request.form.get("projects")
        resume.certifications = request.form.get("certifications")
        resume.languages = request.form.get("languages")
        resume.achievements = request.form.get("achievements")

        resume.linkedin = request.form.get("linkedin")
        resume.github = request.form.get("github")
        resume.portfolio = request.form.get("portfolio")

        db.session.commit()

        return redirect("/dashboard")

    return render_template(
        "edit_resume.html",
        resume=resume
    )


@app.route("/download/<int:id>")
def download_pdf(id):

    resume = Resume.query.get_or_404(id)

    buffer = BytesIO()

    pdf = canvas.Canvas(buffer)

    y = 800

    pdf.setFont("Helvetica-Bold", 18)
    pdf.drawString(50, y, resume.full_name)

    y -= 30
    pdf.setFont("Helvetica", 12)

    pdf.drawString(50, y, f"Email: {resume.email}")
    y -= 20

    pdf.drawString(50, y, f"Phone: {resume.phone}")
    y -= 20

    pdf.drawString(50, y, f"LinkedIn: {resume.linkedin}")
    y -= 20

    pdf.drawString(50, y, f"GitHub: {resume.github}")
    y -= 20

    pdf.drawString(50, y, f"Portfolio: {resume.portfolio}")
    y -= 40

    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(50, y, "Professional Summary")
    y -= 20

    pdf.setFont("Helvetica", 12)
    pdf.drawString(50, y, str(resume.summary))
    y -= 40

    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(50, y, "Education")
    y -= 20

    pdf.setFont("Helvetica", 12)
    pdf.drawString(50, y, str(resume.education))
    y -= 40

    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(50, y, "Skills")
    y -= 20

    pdf.setFont("Helvetica", 12)
    pdf.drawString(50, y, str(resume.skills))
    y -= 40

    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(50, y, "Experience")
    y -= 20

    pdf.setFont("Helvetica", 12)
    pdf.drawString(50, y, str(resume.experience))
    y -= 40

    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(50, y, "Projects")
    y -= 20

    pdf.setFont("Helvetica", 12)
    pdf.drawString(50, y, str(resume.projects))
    y -= 40

    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(50, y, "Certifications")
    y -= 20

    pdf.setFont("Helvetica", 12)
    pdf.drawString(50, y, str(resume.certifications))
    y -= 40

    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(50, y, "Languages")
    y -= 20

    pdf.setFont("Helvetica", 12)
    pdf.drawString(50, y, str(resume.languages))
    y -= 40

    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(50, y, "Achievements")
    y -= 20

    pdf.setFont("Helvetica", 12)
    pdf.drawString(50, y, str(resume.achievements))

    pdf.save()

    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"{resume.full_name}.pdf",
        mimetype="application/pdf"
    )
@app.route("/preview/<int:id>")
def preview_resume(id):

    resume = Resume.query.get_or_404(id)

    return render_template(
        "preview.html",
        resume=resume
    )
@app.route("/download_docx/<int:id>")
def download_docx(id):

    from io import BytesIO

    resume = Resume.query.get_or_404(id)

    document = Document()

    # rest of code...

    document.add_heading(
        resume.full_name,
        level=0
    )

    document.add_paragraph(
        f"Email: {resume.email}"
    )

    document.add_paragraph(
        f"Phone: {resume.phone}"
    )

    document.add_paragraph(
        f"LinkedIn: {resume.linkedin}"
    )

    document.add_paragraph(
        f"GitHub: {resume.github}"
    )

    document.add_paragraph(
        f"Portfolio: {resume.portfolio}"
    )

    document.add_heading(
        "Professional Summary",
        level=1
    )

    document.add_paragraph(
        str(resume.summary)
    )

    document.add_heading(
        "Education",
        level=1
    )

    document.add_paragraph(
        str(resume.education)
    )

    document.add_heading(
        "Skills",
        level=1
    )

    document.add_paragraph(
        str(resume.skills)
    )

    document.add_heading(
        "Experience",
        level=1
    )

    document.add_paragraph(
        str(resume.experience)
    )

    document.add_heading(
        "Projects",
        level=1
    )

    document.add_paragraph(
        str(resume.projects)
    )

    document.add_heading(
        "Certifications",
        level=1
    )

    document.add_paragraph(
        str(resume.certifications)
    )

    document.add_heading(
        "Languages",
        level=1
    )

    document.add_paragraph(
        str(resume.languages)
    )

    document.add_heading(
        "Achievements",
        level=1
    )

    document.add_paragraph(
        str(resume.achievements)
    )
    from io import BytesIO
    buffer = BytesIO()

    document.save(buffer)

    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"{resume.full_name}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
if __name__ == "__main__":
    print("FLASK SERVER STARTING")
    app.run(debug=True)